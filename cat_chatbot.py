import streamlit as st
from docx import Document
import io
import time
import tempfile
import os
from langchain_community.vectorstores import FAISS
# from cat_quiz import show_quiz_sidebar
from rag_core import generate_answer, strip_html, load_and_chunk_document, create_or_update_faiss, retrieve_from_upload_document

# st.set_page_config(page_title="Cat Expert Chatbot 🐾", layout="wide")

def run_chatbot():
    st.markdown("""
        <style>
            .element-container {
                animation: fadeIn 0.3s ease-in;
            }
            @keyframes fadeIn {
                from { opacity: 0; transform: scale(0.95); }
                to { opacity: 1; transform: scale(1); }
            }

            /* Assistance message background color */
            .assistant-message {
            background-color: #fbedf5;
            padding: 12px;
            border-radius: 10px;
            color: black;
            font-size: 12px;
            }
        </style>
    """, unsafe_allow_html=True)

    # Initialize conversation history
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Always show the title once
    st.markdown("""
        <div style='text-align: center;'>
            <h1>😺 Ask the Cat Expert</h1>
            <div style='font-size: 18px; color: #555;'>
                Welcome to the <b>Cat Expert Chatbot Designed by Ming</b>! 🐾<br>
            </div>
            <div style="font-size: 10px; color: #555; font-style: italic;">
                Ask questions about cat health, behavior, grooming, nutrition, plants toxic to cats, or how to raise cats and newborns together — powered by ChatGPT-4o-mini!
            </div>
        </div>
        """, unsafe_allow_html=True)    

    # Add info box on main screen if sidebar might be overlooked:
    st.markdown(
        """
        <div style="text-align: center; font-size: 14px; color: #3B3B3B; background-color: #E1F5FE; 
                    padding: 10px; border-left: 5px solid #2196F3; border-radius: 3px; margin-bottom: 30px;">
            <strong> 🧹Need to start over?</strong> Use the sidebar to reset the conversation.
        </div>
        """,
        unsafe_allow_html=True
    )

    # Reset button right under the title
    with st.sidebar:
        st.header("Options")
        if st.button("🧹 Reset Conversation"):
            st.session_state.messages = []
            st.session_state.reset_done = True
            st.rerun()
        # Download button (simulated with conditional logic)
        if st.button("📥 Download Conversation"):
            if st.session_state.get("messages"):
                # create the word document
                doc = Document()
                doc.add_heading('Conversation History', level = 1)
                for m in st.session_state.messages:
                    role = m['role'].capitalize()
                    content = strip_html(m['content'])
                    doc.add_paragraph(f"{role}: {content}")
                
                # Save to a BytesIO buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Create download button
                st.download_button(
                    label="Click to confirm download",
                    data=buffer,
                    file_name="cat_chat.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download"
                )
                
            else:
                st.warning("⚠️ The conversation is empty. Nothing to download.")


    # Outside sidebar, after page reload show success message after reset
    if st.session_state.get("reset_done", False):
        st.success("Conversation reset. You can start a new conversation")
        del st.session_state["reset_done"]

    # Display chat history (exclude the *last* assistant message if a new one is being typed)
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"], unsafe_allow_html=True)

    # Initialize FAISS vectorstore in session state if not already
    if "faiss_store" not in st.session_state:
        st.session_state.faiss_store = None

    # User input box (chat-style)
    if prompt := st.chat_input("Ask me anything about cats..."):
        # Show user's message
        with st.chat_message("user"):
            st.markdown(prompt)

        # Generate assistant response
        with st.spinner("Thinking..."):
            try:
                # get chat history
                history = st.session_state.messages[-6:] # 3 turns = 6 messages
                # get user-uploaded document
                if "uploaded_docs" in st.session_state:
                    user_doc_retrieval = retrieve_from_upload_document(prompt)
                    # generate response
                    response = generate_answer(prompt, history, user_doc_retrieval)
                else:
                    # generate response
                    response = generate_answer(prompt, history)
            except Exception as e:
                response = f"⚠️ Oops! Something went wrong.\n\n**Error:** `{str(e)}`"
                st.error(f"Exception: {e}")  # Optional: show in sidebar too
                print(f"error: {e}")

        with st.chat_message("assistant"):
            formatted_response = f"""
                <div class="assistant-message">
                    {response}
                </div>
            """
            st.markdown(formatted_response, unsafe_allow_html=True)

        # Save messages to session_state (already formatted!)
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.session_state.messages.append({"role": "assistant", "content": formatted_response})
        # st.rerun() 

    # add feature to allow user's upload of document file
    left_col, right_col = st.columns(2)
    with left_col:
        uploaded_doc = st.file_uploader("📄 Upload Document\nto get answer based on your uploaded document", type=["pdf", "txt", "docx"], key="doc_upload")
        if uploaded_doc:
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_doc.name)[1]) as tmp:
                tmp.write(uploaded_doc.read())
                doc_chunks = load_and_chunk_document(tmp.name)
                st.session_state.faiss_store = create_or_update_faiss(st.session_state.faiss_store, doc_chunks)
                # track state of the uploaded document
                if "uploaded_docs" not in st.session_state:
                    st.session_state.uploaded_docs = []
                st.session_state.uploaded_docs.append(uploaded_doc.name)

    # add feature to allow user's upload of plant or nutrition fact image file
    with right_col:
        uploaded_image = st.file_uploader("📷 Upload Image\nto analyze the ingredients of the cat food", type=["png", "jpg", "jpeg"], key="img_upload")
        if uploaded_image:
            st.image(uploaded_image, caption="Uploaded Image", use_column_width=True)